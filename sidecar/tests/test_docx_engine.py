"""Unit tests for Paragraph-Level DOCX Style Rebuilding Engine."""

import os
import tempfile
import docx
from docx import Document
from docx.shared import Pt, Inches
import pytest
from app.docx_engine import DocxEngine, MAX_PROJECTS, MAX_BULLETS_PER_PROJECT

@pytest.fixture
def docx_fixture():
    tmpdir = tempfile.mkdtemp()
    template_path = os.path.join(tmpdir, "template_resume.docx")
    output_path = os.path.join(tmpdir, "compiled_resume.docx")

    # Create dummy master template
    doc = Document()
    p_header = doc.add_paragraph("John Doe | Full Stack Engineer")
    p_header.paragraph_format.space_after = Pt(6)

    doc.add_paragraph("TECHNICAL SKILLS")
    p_skills = doc.add_paragraph("{{SKILLS}}")
    p_skills.paragraph_format.line_spacing = 1.15

    doc.add_paragraph("PROJECT EXPERIENCE")
    p_projects = doc.add_paragraph("{{PROJECTS}}")
    p_projects.paragraph_format.line_spacing = 1.15

    doc.save(template_path)

    yield {
        "template_path": template_path,
        "output_path": output_path,
        "tmpdir": tmpdir
    }

    if os.path.exists(template_path):
        os.remove(template_path)
    if os.path.exists(output_path):
        os.remove(output_path)
    os.rmdir(tmpdir)

def test_docx_rebuilding_replaces_placeholders(docx_fixture):
    """Assert placeholders {{PROJECTS}} and {{SKILLS}} are replaced cleanly."""
    projects_data = [
        {
            "title": "Distributed KV Engine",
            "tech_stack": "Go, Raft, gRPC",
            "live_demo_url": "https://kv.example.com",
            "bullets": [
                "Implemented Raft consensus protocol handling 10k ops/sec.",
                "Reduced write latency by 45% using memory-mapped SSTables."
            ]
        },
        {
            "title": "Cloud Analytics Dashboard",
            "tech_stack": "React, TypeScript, FastAPI",
            "live_demo_url": "https://analytics.example.com",
            "bullets": [
                "Engineered real-time telemetry visualizer with SSE streams."
            ]
        }
    ]

    skills_data = {
        "Languages": ["Go", "Python", "TypeScript", "Rust"],
        "Frameworks": ["FastAPI", "React", "Node.js"]
    }

    out_file = DocxEngine.rebuild_resume(
        template_path=docx_fixture["template_path"],
        output_path=docx_fixture["output_path"],
        projects=projects_data,
        skills=skills_data
    )

    assert os.path.exists(out_file)

    # Inspect generated document
    rebuilt_doc = Document(out_file)
    all_text = "\n".join(p.text for p in rebuilt_doc.paragraphs)

    assert "{{PROJECTS}}" not in all_text
    assert "{{SKILLS}}" not in all_text
    assert "Distributed KV Engine" in all_text
    assert "Cloud Analytics Dashboard" in all_text
    assert "Implemented Raft consensus protocol handling 10k ops/sec." in all_text
    assert "Languages: Go, Python, TypeScript, Rust" in all_text

    # Verify hyperlink relationships exist
    rel_targets = [rel.target_ref for rel in rebuilt_doc.part.rels.values() if rel.is_external]
    assert "https://kv.example.com" in rel_targets
    assert "https://analytics.example.com" in rel_targets

def test_single_page_guardrail_capping(docx_fixture):
    """Assert engine enforces max 4 projects and max 4 bullets per project."""
    # Create 7 projects, each with 6 bullets
    oversized_projects = []
    for i in range(1, 8):
        oversized_projects.append({
            "title": f"Project {i}",
            "tech_stack": "Tech",
            "bullets": [f"Bullet {j} for project {i}" for j in range(1, 7)]
        })

    skills_data = ["Python", "Go", "Docker", "Kubernetes"]

    out_file = DocxEngine.rebuild_resume(
        template_path=docx_fixture["template_path"],
        output_path=docx_fixture["output_path"],
        projects=oversized_projects,
        skills=skills_data
    )

    rebuilt_doc = Document(out_file)
    all_text = "\n".join(p.text for p in rebuilt_doc.paragraphs)

    # Projects 1-4 must exist
    for i in range(1, 5):
        assert f"Project {i}" in all_text
        # Bullets 1-4 must exist
        for j in range(1, 5):
            assert f"Bullet {j} for project {i}" in all_text
        # Bullets 5-6 must NOT exist
        assert f"Bullet 5 for project {i}" not in all_text
        assert f"Bullet 6 for project {i}" not in all_text

    # Projects 5-7 must NOT exist
    for i in range(5, 8):
        assert f"Project {i}" not in all_text
