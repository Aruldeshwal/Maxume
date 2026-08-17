"""Paragraph-Level DOCX Style Cloning & Rebuilding Engine for Maxume."""

import os
import sys
import re
import time
import shutil
from typing import List, Dict, Any, Optional, Union
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
import docx.opc.constants
import docx.oxml.shared

MAX_PROJECTS = 3
MAX_BULLETS_PER_PROJECT = 3

METADATA_FILTER_KEYWORDS = [
    "github", "language", "live demo", "demo", "url", "tech stack", "technology", "status", "http", "https", "repository"
]

def is_valid_bullet_point(text: str) -> bool:
    """Filters out metadata labels, URLs, and headers to ensure only real achievements enter resume."""
    clean = re.sub(r'^[-*•\d.)\s]+', '', text).replace('**', '').replace('__', '').strip()
    if len(clean) < 15 or clean.startswith("#"):
        return False
    lower = clean.lower()
    for kw in METADATA_FILTER_KEYWORDS:
        if lower.startswith(kw + ":") or lower.startswith(kw + " :") or lower.startswith(kw + " /") or lower.startswith("http"):
            return False
    return True

def clean_bullet_string(text: str) -> str:
    """Strips leading markdown symbols, bolding, and bullets."""
    clean = re.sub(r'^[-*•\d.)\s]+', '', text).strip()
    return clean.replace("**", "").replace("__", "")

def add_hyperlink(paragraph, url: str, text: str, color: str = "990000", underline: bool = True):
    """
    Embed an active clickable hyperlink in a python-docx paragraph using precise OXML ordering.
    Verbatim implementation from projectoverview.md §3.
    """
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    if color:
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)
    if underline:
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'single')
        rPr.append(u)
    new_run.append(rPr)
    text_node = docx.oxml.shared.OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def copy_paragraph_style(source_para, target_para):
    """Clone layout settings, indentation, and alignment from a reference paragraph."""
    if source_para.style:
        try:
            target_para.style = source_para.style
        except Exception:
            pass
    target_para.paragraph_format.alignment = source_para.paragraph_format.alignment
    target_para.paragraph_format.left_indent = source_para.paragraph_format.left_indent
    target_para.paragraph_format.right_indent = source_para.paragraph_format.right_indent
    target_para.paragraph_format.space_before = Pt(0)
    target_para.paragraph_format.space_after = Pt(1.5)
    target_para.paragraph_format.line_spacing = 1.05

def insert_paragraph_before(target_para, text="", style=None):
    """Insert a new paragraph before the specified target paragraph in OXML tree."""
    new_p = docx.oxml.shared.OxmlElement('w:p')
    target_para._p.addprevious(new_p)
    new_para = docx.text.paragraph.Paragraph(new_p, target_para._parent)
    if style:
        new_para.style = style
    elif target_para.style:
        new_para.style = target_para.style
    if text:
        new_para.add_run(text)
    return new_para

def remove_paragraph(paragraph):
    """Safely remove a paragraph element from its parent XML element."""
    p_element = paragraph._p
    parent = p_element.getparent()
    if parent is not None:
        parent.remove(p_element)

def create_default_master_template(target_path: str):
    """Generates a default professional master resume docx with standard placeholders."""
    os.makedirs(os.path.dirname(os.path.abspath(target_path)), exist_ok=True)
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
    p_name = doc.add_paragraph()
    r_name = p_name.add_run("ARUL DESHWAL")
    r_name.bold = True
    r_name.font.size = Pt(16)
    p_name.paragraph_format.space_after = Pt(1)
    
    p_contact = doc.add_paragraph("Sector 84, Faridabad, India  |  +91 80763 55239  |  aruldeshwal1@gmail.com  |  LinkedIn  |  GitHub")
    p_contact.paragraph_format.space_after = Pt(4)
    
    p_sum_head = doc.add_paragraph()
    r_sh = p_sum_head.add_run("SUMMARY")
    r_sh.bold = True
    p_sum_head.paragraph_format.space_before = Pt(3)
    p_sum_head.paragraph_format.space_after = Pt(1)
    
    p_sum = doc.add_paragraph("CS undergraduate (8.6 CGPA) seeking SDE/MERN intern roles. Experienced in building full-stack MERN/Next.js applications, designing RESTful APIs, and structuring PostgreSQL/MongoDB databases.")
    p_sum.paragraph_format.space_after = Pt(4)
    
    p_sk_head = doc.add_paragraph()
    r_skh = p_sk_head.add_run("CORE SKILLS")
    r_skh.bold = True
    p_sk_head.paragraph_format.space_before = Pt(3)
    p_sk_head.paragraph_format.space_after = Pt(1)
    
    doc.add_paragraph("{{SKILLS}}")
    
    p_pr_head = doc.add_paragraph()
    r_prh = p_pr_head.add_run("PROJECTS")
    r_prh.bold = True
    p_pr_head.paragraph_format.space_before = Pt(3)
    p_pr_head.paragraph_format.space_after = Pt(1)
    
    doc.add_paragraph("{{PROJECTS}}")
    
    p_ed_head = doc.add_paragraph()
    r_edh = p_ed_head.add_run("EDUCATION")
    r_edh.bold = True
    p_ed_head.paragraph_format.space_before = Pt(3)
    p_ed_head.paragraph_format.space_after = Pt(1)
    
    doc.add_paragraph("Bachelor of Technology (B.Tech), Computer Science & Engineering\nADGITM, New Delhi, India  |  CGPA: 8.6  |  Sep 2023 – Jul 2027 (Expected)")
    
    doc.save(target_path)
    return target_path

def resolve_master_template(preferred_path: Optional[str] = None) -> str:
    """Discovers Master_Resume.docx across all known workspace, desktop, and AppData paths."""
    candidates = []
    if preferred_path:
        candidates.append(preferred_path)
        candidates.append(os.path.abspath(preferred_path))
    
    env_path = os.environ.get("MASTER_RESUME_PATH")
    if env_path:
        candidates.append(env_path)
        candidates.append(os.path.abspath(env_path))
        
    home = os.path.expanduser("~")
    appdata = os.environ.get("APPDATA", "")
    
    candidates.extend([
        os.path.abspath("Master_Resume.docx"),
        os.path.join(os.getcwd(), "Master_Resume.docx"),
        os.path.join(home, "OneDrive", "Desktop", "Maxume", "Master_Resume.docx"),
        os.path.join(home, "Desktop", "Maxume", "Master_Resume.docx"),
        os.path.join(home, "OneDrive", "Desktop", "Master_Resume.docx"),
        os.path.join(home, "Desktop", "Master_Resume.docx"),
        os.path.join(home, "Documents", "Master_Resume.docx"),
        os.path.join(appdata, "Maxume", "Master_Resume.docx") if appdata else "",
        os.path.join(os.path.dirname(sys.executable), "Master_Resume.docx"),
        os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "..", "Master_Resume.docx"),
    ])
    
    for c in candidates:
        if c and os.path.exists(c):
            # Cache to AppData for permanent offline availability in installed mode
            if appdata and not c.startswith(appdata):
                try:
                    appdata_target = os.path.join(appdata, "Maxume", "Master_Resume.docx")
                    os.makedirs(os.path.dirname(appdata_target), exist_ok=True)
                    if not os.path.exists(appdata_target):
                        shutil.copy2(c, appdata_target)
                except Exception:
                    pass
            return os.path.abspath(c)
            
    # Auto-generate if missing
    target_dir = os.path.join(appdata, "Maxume") if appdata else os.getcwd()
    fallback_template = os.path.join(target_dir, "Master_Resume.docx")
    create_default_master_template(fallback_template)
    return fallback_template

class DocxEngine:
    @staticmethod
    def rebuild_resume(
        template_path: Optional[str],
        output_path: str,
        projects: List[Dict[str, Any]],
        skills: Union[List[str], Dict[str, List[str]]],
        hyperlink_color: str = "990000"
    ) -> str:
        """
        Rebuilds a resume document by replacing {{PROJECTS}} and {{SKILLS}} placeholders
        with styled content, embedding tech stack and timeline, and maximizing high-impact bullet coverage.
        """
        actual_template = template_path if (template_path and os.path.exists(template_path)) else resolve_master_template(template_path)
        doc = Document(actual_template)
        
        # Enforce single-page guardrail: 3 projects max, 3-4 bullets each to fill available lines
        bounded_projects = projects[:MAX_PROJECTS]
        max_bullets_per_proj = 4 if len(bounded_projects) <= 2 else 3
        
        # Locate placeholder paragraphs
        projects_placeholder = None
        skills_placeholder = None

        for p in doc.paragraphs:
            if "{{PROJECTS}}" in p.text:
                projects_placeholder = p
            if "{{SKILLS}}" in p.text:
                skills_placeholder = p

        # 1. Inject Projects
        if projects_placeholder is not None:
            for proj in bounded_projects:
                title = proj.get("title") or proj.get("directory_name") or "Project"
                tech_stack = proj.get("tech_stack", "")
                demo_url = proj.get("live_demo_url") or proj.get("url")
                date_str = proj.get("date") or proj.get("timeline", "")
                
                # Project Title Heading Paragraph
                title_para = insert_paragraph_before(projects_placeholder)
                copy_paragraph_style(projects_placeholder, title_para)
                title_para.paragraph_format.space_before = Pt(3)
                title_para.paragraph_format.space_after = Pt(1)
                title_para.paragraph_format.line_spacing = 1.0
                
                # Add hyperlink if live demo URL exists, otherwise plain bold title
                if demo_url:
                    add_hyperlink(title_para, demo_url, title, color=hyperlink_color, underline=True)
                else:
                    r = title_para.add_run(title)
                    r.bold = True
                
                # Add tech stack & timeline to title line in brief
                if tech_stack and "general" not in tech_stack.lower():
                    r_tech = title_para.add_run(f" | {tech_stack}")
                    r_tech.italic = True

                if date_str:
                    r_date = title_para.add_run(f" | {date_str}")
                    r_date.italic = True

                # Filter and inject authentic engineering bullet points
                bullets = proj.get("bullets", [])
                clean_bullets = []
                for b in bullets:
                    if is_valid_bullet_point(b):
                        clean_bullets.append(clean_bullet_string(b))

                if not clean_bullets:
                    summary_text = proj.get("summary_markdown", "")
                    if summary_text:
                        for line in summary_text.splitlines():
                            if is_valid_bullet_point(line):
                                clean_bullets.append(clean_bullet_string(line))

                for bullet_text in clean_bullets[:max_bullets_per_proj]:
                    bullet_para = insert_paragraph_before(projects_placeholder)
                    copy_paragraph_style(projects_placeholder, bullet_para)
                    bullet_para.paragraph_format.left_indent = Inches(0.25)
                    bullet_para.paragraph_format.space_before = Pt(0)
                    bullet_para.paragraph_format.space_after = Pt(1.5)
                    bullet_para.paragraph_format.line_spacing = 1.05
                    bullet_para.add_run(f"•  {bullet_text}")

            # Remove original {{PROJECTS}} placeholder
            remove_paragraph(projects_placeholder)

        # 2. Inject Skills
        if skills_placeholder is not None:
            if isinstance(skills, dict):
                for category, skill_list in skills.items():
                    skill_para = insert_paragraph_before(skills_placeholder)
                    copy_paragraph_style(skills_placeholder, skill_para)
                    skill_para.paragraph_format.space_before = Pt(0)
                    skill_para.paragraph_format.space_after = Pt(1.5)
                    r_cat = skill_para.add_run(f"{category}: ")
                    r_cat.bold = True
                    skill_para.add_run(", ".join(skill_list))
            elif isinstance(skills, list):
                skill_para = insert_paragraph_before(skills_placeholder)
                copy_paragraph_style(skills_placeholder, skill_para)
                skill_para.paragraph_format.space_before = Pt(0)
                skill_para.paragraph_format.space_after = Pt(1.5)
                skill_para.add_run(" • ".join(skills))
            
            # Remove original {{SKILLS}} placeholder
            remove_paragraph(skills_placeholder)

        # Save to output path with fallback if file is currently open/locked in Word
        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
        try:
            doc.save(output_path)
            return output_path
        except (PermissionError, IOError):
            base_dir = os.path.dirname(os.path.abspath(output_path))
            file_name, ext = os.path.splitext(os.path.basename(output_path))
            fallback_path = os.path.join(base_dir, f"{file_name}_new{ext}")
            try:
                doc.save(fallback_path)
                return fallback_path
            except (PermissionError, IOError):
                timestamp_path = os.path.join(base_dir, f"{file_name}_{int(time.time())}{ext}")
                doc.save(timestamp_path)
                return timestamp_path
