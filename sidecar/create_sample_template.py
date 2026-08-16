"""Generate sample Master_Resume.docx template with placeholders for Maxume."""

import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_sample_master_resume(target_path: str = "Master_Resume.docx"):
    doc = Document()
    
    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # Name Header
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_name = p_name.add_run("ALEXANDER MERCER")
    r_name.bold = True
    r_name.font.size = Pt(18)
    r_name.font.color.rgb = RGBColor(190, 18, 60) # Legion Crimson

    # Contact Details Line (PII - Stays Local)
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_after = Pt(12)
    p_contact.add_run("San Francisco, CA • (555) 234-5678 • alex.mercer@example.com • ")
    r_li = p_contact.add_run("linkedin.com/in/alex-mercer")
    r_li.font.color.rgb = RGBColor(100, 116, 139)

    # Section 1: Technical Skills
    p_skills_hdr = doc.add_paragraph()
    p_skills_hdr.paragraph_format.space_before = Pt(8)
    p_skills_hdr.paragraph_format.space_after = Pt(2)
    r_sh = p_skills_hdr.add_run("TECHNICAL SKILLS")
    r_sh.bold = True
    r_sh.font.size = Pt(11)
    
    # Skills placeholder paragraph
    p_skills_ph = doc.add_paragraph("{{SKILLS}}")
    p_skills_ph.paragraph_format.space_after = Pt(8)
    p_skills_ph.paragraph_format.line_spacing = 1.15

    # Section 2: Experience
    p_exp_hdr = doc.add_paragraph()
    p_exp_hdr.paragraph_format.space_before = Pt(8)
    p_exp_hdr.paragraph_format.space_after = Pt(2)
    r_eh = p_exp_hdr.add_run("WORK EXPERIENCE")
    r_eh.bold = True
    r_eh.font.size = Pt(11)

    p_job1 = doc.add_paragraph()
    r_j1 = p_job1.add_run("Senior Software Engineer")
    r_j1.bold = True
    p_job1.add_run(" | CloudScale Systems • San Francisco, CA (2022 – Present)")
    p_job1_b1 = doc.add_paragraph("• Architected low-latency ingestion pipelines processing 2.5B events daily with Apache Kafka and Go.")
    p_job1_b1.paragraph_format.left_indent = Inches(0.2)
    p_job1_b2 = doc.add_paragraph("• Led migration of core monolithic database to distributed TiDB cluster, cutting p99 query latency by 40%.")
    p_job1_b2.paragraph_format.left_indent = Inches(0.2)

    # Section 3: Key Technical Projects
    p_proj_hdr = doc.add_paragraph()
    p_proj_hdr.paragraph_format.space_before = Pt(10)
    p_proj_hdr.paragraph_format.space_after = Pt(2)
    r_ph = p_proj_hdr.add_run("KEY TECHNICAL PROJECTS")
    r_ph.bold = True
    r_ph.font.size = Pt(11)

    # Projects placeholder paragraph
    p_proj_ph = doc.add_paragraph("{{PROJECTS}}")
    p_proj_ph.paragraph_format.space_after = Pt(8)
    p_proj_ph.paragraph_format.line_spacing = 1.15

    # Section 4: Education
    p_edu_hdr = doc.add_paragraph()
    p_edu_hdr.paragraph_format.space_before = Pt(8)
    p_edu_hdr.paragraph_format.space_after = Pt(2)
    r_ed = p_edu_hdr.add_run("EDUCATION")
    r_ed.bold = True
    r_ed.font.size = Pt(11)

    p_edu = doc.add_paragraph()
    r_deg = p_edu.add_run("B.S. in Computer Science")
    r_deg.bold = True
    p_edu.add_run(" | University of California, Berkeley (2018 – 2022)")

    doc.save(target_path)
    print(f"Master Resume template generated at: {target_path}")

if __name__ == "__main__":
    create_sample_master_resume("Master_Resume.docx")
